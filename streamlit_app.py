
import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
import io
import json
import os
import pandas as pd
from docx import Document
from PIL import Image

# Configuration
st.set_page_config(page_title="SmartPDF Converter", page_icon="📄", layout="centered")

# CSS for styling
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stButton>button { width: 100%; border-radius: 12px; height: 3em; font-weight: bold; }
    .stDownloadButton>button { width: 100%; border-radius: 12px; background-color: #059669; color: white; }
    </style>
    """, unsafe_allow_value=True)

def pdf_to_images(pdf_file):
    """Converts first 5 pages of PDF to images."""
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    images = []
    # Limit to 5 pages for performance and token management
    for i in range(min(len(doc), 5)):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_data = pix.tobytes("jpeg")
        img = Image.open(io.BytesIO(img_data))
        images.append(img)
    doc.close()
    return images

def analyze_document(images, output_format):
    """Sends images to Gemini for analysis."""
    api_key = os.environ.get("API_KEY")
    if not api_key:
        st.error("API_KEY not found in environment variables.")
        return None

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash') # Using stable flash for Python SDK

    if output_format == "Word":
        prompt = "Analyze the provided PDF pages. Reconstruct the document content in structured Markdown. Preserve headings, lists, and bold text. Return a JSON object with a 'content' field."
    else:
        prompt = "Analyze the provided PDF pages. Extract all tabular data. Return a JSON object with a 'tables' field (a 2D array of strings)."

    # Convert PIL images to bytes for the API
    parts = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format='JPEG')
        parts.append({"mime_type": "image/jpeg", "data": buf.getvalue()})
    
    parts.append(prompt)
    
    response = model.generate_content(
        parts,
        generation_config={"response_mime_type": "application/json"}
    )
    
    return json.loads(response.text)

def create_docx(content):
    """Generates a .docx file from markdown content."""
    doc = Document()
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        else:
            doc.add_paragraph(line)
    
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def create_excel(tables):
    """Generates an .xlsx file from 2D list."""
    df = pd.DataFrame(tables)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, header=False)
    return output.getvalue()

def main():
    st.title("📄 SmartPDF Converter")
    st.subheader("AI-Powered PDF to Word & Excel")
    
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
    if uploaded_file:
        col1, col2 = st.columns(2)
        with col1:
            output_format = st.radio("Convert to:", ["Word", "Excel"])
        
        if st.button("🚀 Convert Now"):
            try:
                with st.status("Processing document...", expanded=True) as status:
                    st.write("Extracting pages...")
                    images = pdf_to_images(uploaded_file)
                    
                    st.write("AI Analysis (this may take a moment)...")
                    result = analyze_document(images, output_format)
                    
                    st.write("Generating file...")
                    if output_format == "Word":
                        file_bytes = create_docx(result.get('content', ''))
                        ext = "docx"
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        file_bytes = create_excel(result.get('tables', []))
                        ext = "xlsx"
                        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    
                    status.update(label="Conversion Complete!", state="complete")

                st.success(f"Successfully converted to {output_format}!")
                st.download_button(
                    label=f"⬇️ Download {output_format} File",
                    data=file_bytes,
                    file_name=f"converted_document.{ext}",
                    mime=mime
                )
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    st.divider()
    st.caption("Note: For this demo, conversion is limited to the first 5 pages. Powered by Gemini AI.")

if __name__ == "__main__":
    main()
